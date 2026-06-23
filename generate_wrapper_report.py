#!/usr/bin/env python3
"""Generate a shareable wrapper validation report.

Outputs a self-contained folder with:
- HTML report with playable videos,
- DOCX report with figures, tables, and video paths,
- copied figures/videos/CSV,
- ZIP archive for sending to collaborators.
"""

from __future__ import annotations

import csv
import html
import os
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
FIG_DIR = ROOT / "figures"
COMPARE_DIR = FIG_DIR / "cost_profile_comparison"
REPORT_DIR = ROOT / "reports" / "wrapper_validation_report"
ASSET_DIR = REPORT_DIR / "assets"

SUMMARY_CSV = COMPARE_DIR / "summary.csv"
FIGURES = {
    "safety_table": COMPARE_DIR / "overview_safety_table.png",
    "metrics": COMPARE_DIR / "overview_metrics.png",
    "trajectories": COMPARE_DIR / "overview_trajectories.png",
}
VIDEOS = {
    "safe": FIG_DIR / "mgigo_borrow_overtake_safe.mp4",
    "blocked": FIG_DIR / "mgigo_borrow_overtake_blocked.mp4",
    "critical": FIG_DIR / "mgigo_borrow_overtake_critical.mp4",
}


def _clean_report_dir() -> None:
    if REPORT_DIR.exists():
        shutil.rmtree(REPORT_DIR)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def _copy_assets() -> dict[str, str]:
    copied = {}
    for name, path in FIGURES.items():
        dst = ASSET_DIR / path.name
        shutil.copy2(path, dst)
        copied[name] = f"assets/{dst.name}"
    for name, path in VIDEOS.items():
        dst = ASSET_DIR / path.name
        shutil.copy2(path, dst)
        copied[f"video_{name}"] = f"assets/{dst.name}"
    dst_csv = ASSET_DIR / SUMMARY_CSV.name
    shutil.copy2(SUMMARY_CSV, dst_csv)
    copied["summary_csv"] = f"assets/{dst_csv.name}"
    return copied


def _load_summary() -> list[dict[str, str]]:
    with SUMMARY_CSV.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def _short_cost(cost: str) -> str:
    if cost == "borrow_overtake":
        return "wrapper"
    if cost == "borrow_overtake_baseline":
        return "old baseline"
    if cost == "borrow_overtake_matched":
        return "matched baseline"
    return cost


def _short_scenario(scenario: str) -> str:
    return scenario.replace("borrow_overtake_", "")


def _fmt_float(value: str, digits: int = 2) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return value


def _summary_rows(rows: list[dict[str, str]]) -> list[list[str]]:
    result = []
    for row in rows:
        result.append([
            _short_scenario(row["scenario"]),
            _short_cost(row["cost_profile"]),
            row["pass_success"],
            row["returned_to_lane"],
            row["conflict_free_while_borrowing"],
            _fmt_float(row["min_gap_while_borrowing"], 1),
            _fmt_float(row["min_ttc_while_borrowing"], 2),
            _fmt_float(row["final_y"], 2),
            _fmt_float(row["final_v"], 1),
            _fmt_float(row["final_lead"], 1),
        ])
    return result


def _html_table(headers: list[str], rows: list[list[str]]) -> str:
    head = "".join(f"<th>{html.escape(h)}</th>" for h in headers)
    body_rows = []
    for row in rows:
        cells = "".join(f"<td>{html.escape(str(cell))}</td>" for cell in row)
        body_rows.append(f"<tr>{cells}</tr>")
    return f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"


def build_html(asset_paths: dict[str, str], rows: list[dict[str, str]]) -> Path:
    headers = [
        "scenario",
        "cost",
        "pass+return",
        "returned",
        "conflict-free while borrowing",
        "min gap while borrowing (m)",
        "min TTC while borrowing (s)",
        "final y (m)",
        "final v (m/s)",
        "final lead (m)",
    ]
    table = _html_table(headers, _summary_rows(rows))
    generated = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    html_text = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>Constraint Wrapper Validation Report</title>
  <style>
    body {{ margin: 0; font-family: Arial, "Microsoft YaHei", sans-serif; background: #0a1120; color: #edf3ff; }}
    main {{ max-width: 1180px; margin: 0 auto; padding: 34px 28px 56px; }}
    h1 {{ font-size: 30px; margin: 0 0 8px; }}
    h2 {{ margin-top: 34px; border-bottom: 1px solid #33445f; padding-bottom: 8px; }}
    p, li {{ line-height: 1.68; color: #d9e3f4; }}
    code {{ color: #9ff3bd; }}
    .card {{ background: #10192b; border: 1px solid #30425f; border-radius: 8px; padding: 18px; margin: 16px 0; }}
    table {{ width: 100%; border-collapse: collapse; margin: 14px 0; font-size: 14px; }}
    th, td {{ border: 1px solid #40506b; padding: 9px 8px; text-align: center; }}
    th {{ background: #1b2b46; color: white; }}
    tr:nth-child(even) td {{ background: #132037; }}
    tr:nth-child(odd) td {{ background: #10192b; }}
    img {{ width: 100%; border: 1px solid #30425f; border-radius: 6px; background: #0a1120; }}
    .videos {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 14px; }}
    video {{ width: 100%; border: 1px solid #30425f; border-radius: 6px; background: black; }}
    .muted {{ color: #9fb0c8; font-size: 13px; }}
  </style>
</head>
<body>
<main>
  <h1>Constraint Wrapper Validation Report</h1>
  <p class="muted">Generated: {generated}</p>

  <div class="card">
    <p><b>结论摘要：</b>matched baseline 不调用 wrapper API，但手写复现 wrapper 的 log/saturation/hard/tunable/priority 变换。测试中 matched 与 wrapper 行为和指标一致，说明效果来自这套 constraint-to-cost transformation；wrapper 的优势是把这套变换结构化、参数化、可复用，降低手写层级和饱和顺序出错风险。</p>
  </div>

  <h2>1. 测试设置</h2>
  <ul>
    <li>场景：<code>borrow_overtake_safe</code>, <code>borrow_overtake_blocked</code>, <code>borrow_overtake_critical</code></li>
    <li>求解器：同一 MGIGO / MPC 闭环求解器，同一随机种子。</li>
    <li>底层项：三组 cost 使用同一批 objective / STL robustness violation 函数。</li>
    <li>判据：return tolerance = 0.6 m；借道期间 gap >= 28 m 或 TTC >= 1.5 s 判为 conflict-free。</li>
  </ul>

  <h2>2. Cost Profile 对比</h2>
  <table>
    <thead><tr><th>profile</th><th>说明</th><th>用途</th></tr></thead>
    <tbody>
      <tr><td>wrapper</td><td>调用 <code>constraint_dsl.build()</code>，声明 hard/tunable/priority。</td><td>同事 wrapper 主方案</td></tr>
      <tr><td>old baseline</td><td>旧手写 hierarchical cost，饱和尺度和 hard penalty 与 wrapper 不一致。</td><td>对比旧手写方法</td></tr>
      <tr><td>matched baseline</td><td>不调用 wrapper，但手写复现 wrapper 的数学变换。</td><td>控制变量验证</td></tr>
    </tbody>
  </table>

  <h2>3. 结果表</h2>
  {table}
  <p>关键现象：critical 场景中 old baseline 完成超车，但借道期间 min gap 和 min TTC 为负，表示对向冲突；wrapper 与 matched baseline 均选择安全等待/放弃。</p>

  <h2>4. 总览图</h2>
  <h3>安全判据表</h3>
  <img src="{asset_paths['safety_table']}" alt="safety table">
  <h3>指标曲线</h3>
  <img src="{asset_paths['metrics']}" alt="metrics overview">
  <h3>轨迹总览</h3>
  <img src="{asset_paths['trajectories']}" alt="trajectory overview">

  <h2>5. 视频</h2>
  <div class="videos">
    <div><h3>safe</h3><video controls src="{asset_paths['video_safe']}"></video></div>
    <div><h3>blocked</h3><video controls src="{asset_paths['video_blocked']}"></video></div>
    <div><h3>critical</h3><video controls src="{asset_paths['video_critical']}"></video></div>
  </div>

  <h2>6. 可复现实验命令</h2>
  <pre><code>cd /mnt/d/claude_workspace1/IOC_AGV-main/IOC_AGV-main
JAX_PLATFORMS=cuda uv run python compare_cost_profiles.py --force</code></pre>
</main>
</body>
</html>
"""
    path = REPORT_DIR / "wrapper_validation_report.html"
    path.write_text(html_text, encoding="utf-8")
    return path


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_image(doc: Document, path: Path, width: float = 6.7) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))


def build_docx(rows: list[dict[str, str]]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("Constraint Wrapper Validation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("MGIGO borrow-lane overtaking cost transformation A/B test")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    _add_heading(doc, "1. 结论摘要", 1)
    doc.add_paragraph(
        "matched baseline 不调用 wrapper API，但手写复现 wrapper 的 log/saturation/"
        "hard/tunable/priority 变换。测试中 matched 与 wrapper 行为和指标一致，"
        "说明效果来自这套 constraint-to-cost transformation；wrapper 的实际优势是"
        "结构化、参数化、可复用，降低手写层级和饱和顺序出错风险。"
    )

    _add_heading(doc, "2. 测试配置", 1)
    for item in [
        "场景：borrow_overtake_safe, borrow_overtake_blocked, borrow_overtake_critical。",
        "求解器：同一 MGIGO / MPC 闭环求解器，同一随机种子。",
        "底层项：三组 cost 使用同一批 objective / STL robustness violation 函数。",
        "判据：return tolerance = 0.6 m；借道期间 gap >= 28 m 或 TTC >= 1.5 s 判为 conflict-free。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_heading(doc, "3. Cost Profile 差异", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "profile"
    hdr[1].text = "说明"
    hdr[2].text = "用途"
    for profile, desc, use in [
        ("wrapper", "调用 constraint_dsl.build()，声明 hard/tunable/priority。", "同事 wrapper 主方案"),
        ("old baseline", "旧手写 hierarchical cost，饱和尺度和 hard penalty 与 wrapper 不一致。", "对比旧手写方法"),
        ("matched baseline", "不调用 wrapper，但手写复现 wrapper 的数学变换。", "控制变量验证"),
    ]:
        cells = table.add_row().cells
        cells[0].text = profile
        cells[1].text = desc
        cells[2].text = use

    _add_heading(doc, "4. 结果表", 1)
    headers = [
        "scenario",
        "cost",
        "pass+return",
        "returned",
        "conflict-free",
        "min gap",
        "min TTC",
        "final y",
        "final v",
        "final lead",
    ]
    result_table = doc.add_table(rows=1, cols=len(headers))
    result_table.style = "Table Grid"
    for idx, h in enumerate(headers):
        result_table.rows[0].cells[idx].text = h
    for row in _summary_rows(rows):
        cells = result_table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)

    doc.add_paragraph(
        "关键现象：critical 场景中 old baseline 完成超车，但借道期间 min gap 和 "
        "min TTC 为负，表示对向冲突；wrapper 与 matched baseline 均选择安全等待/放弃。"
    )

    _add_heading(doc, "5. 总览图", 1)
    for title, path in [
        ("安全判据表", FIGURES["safety_table"]),
        ("指标曲线", FIGURES["metrics"]),
        ("轨迹总览", FIGURES["trajectories"]),
    ]:
        doc.add_heading(title, level=2)
        _add_image(doc, path)

    _add_heading(doc, "6. 视频文件", 1)
    doc.add_paragraph("HTML 报告中可直接播放视频；DOCX 中列出视频路径：")
    for name, path in VIDEOS.items():
        doc.add_paragraph(f"{name}: {path}", style="List Bullet")

    _add_heading(doc, "7. 可复现实验命令", 1)
    doc.add_paragraph(
        "cd /mnt/d/claude_workspace1/IOC_AGV-main/IOC_AGV-main\n"
        "JAX_PLATFORMS=cuda uv run python compare_cost_profiles.py --force"
    )

    path = REPORT_DIR / "wrapper_validation_report.docx"
    doc.save(path)
    return path


def build_zip() -> Path:
    zip_path = REPORT_DIR.parent / "wrapper_validation_report.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in REPORT_DIR.rglob("*"):
            zf.write(path, path.relative_to(REPORT_DIR.parent))
    return zip_path


def main() -> None:
    required = [SUMMARY_CSV, *FIGURES.values(), *VIDEOS.values()]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required report inputs:\n" + "\n".join(missing))

    _clean_report_dir()
    asset_paths = _copy_assets()
    rows = _load_summary()
    html_path = build_html(asset_paths, rows)
    docx_path = build_docx(rows)
    zip_path = build_zip()
    print(f"[save] {html_path}")
    print(f"[save] {docx_path}")
    print(f"[save] {zip_path}")


if __name__ == "__main__":
    main()
